#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Módulo de Análise de Metadados de Arquivos para Ferramenta OSINT
Este módulo permite extrair e analisar metadados de diferentes tipos de arquivos.
"""

import os
import json
import logging
import exifread
import datetime
import pandas as pd
from PIL import Image
from PIL.ExifTags import TAGS, GPSTAGS
from typing import Dict, List, Any, Optional, Union, Tuple
from docx import Document
from openpyxl import load_workbook
from PyPDF2 import PdfReader

# Configuração de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger('osint_metadata_analysis')

class MetadataAnalysisOSINT:
    """Classe principal para análise de metadados de arquivos."""
    
    def __init__(self, output_dir: str = "resultados"):
        """
        Inicializa o módulo de análise de metadados.
        
        Args:
            output_dir: Diretório para salvar os resultados
        """
        self.output_dir = output_dir
        self._setup_directories()
        self.results = {}
        logger.info("Módulo de análise de metadados inicializado")
    
    def _setup_directories(self):
        """Cria os diretórios necessários para armazenar os resultados."""
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            logger.info(f"Diretório de resultados criado: {self.output_dir}")
        
        # Diretórios específicos para cada tipo de arquivo
        for file_type in ['images', 'documents', 'spreadsheets', 'pdfs', 'others']:
            type_dir = os.path.join(self.output_dir, file_type)
            if not os.path.exists(type_dir):
                os.makedirs(type_dir)
                logger.info(f"Diretório para {file_type} criado: {type_dir}")
    
    def analyze_file(self, file_path: str) -> Dict[str, Any]:
        """
        Analisa os metadados de um arquivo com base em sua extensão.
        
        Args:
            file_path: Caminho para o arquivo
            
        Returns:
            Dicionário com metadados do arquivo
        """
        try:
            logger.info(f"Analisando metadados do arquivo: {file_path}")
            
            # Verificar se o arquivo existe
            if not os.path.exists(file_path):
                logger.error(f"Arquivo não encontrado: {file_path}")
                return {"error": f"Arquivo não encontrado: {file_path}"}
            
            # Obter extensão do arquivo
            _, ext = os.path.splitext(file_path)
            ext = ext.lower()
            
            # Direcionar para o método apropriado com base na extensão
            if ext in ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.gif', '.bmp']:
                metadata = self.analyze_image(file_path)
                file_type = 'images'
            elif ext in ['.doc', '.docx']:
                metadata = self.analyze_document(file_path)
                file_type = 'documents'
            elif ext in ['.xls', '.xlsx', '.csv']:
                metadata = self.analyze_spreadsheet(file_path)
                file_type = 'spreadsheets'
            elif ext in ['.pdf']:
                metadata = self.analyze_pdf(file_path)
                file_type = 'pdfs'
            else:
                metadata = self.analyze_generic_file(file_path)
                file_type = 'others'
            
            # Adicionar informações básicas do arquivo
            file_info = self._get_file_info(file_path)
            metadata.update(file_info)
            
            # Salvar resultados
            self._save_results(file_type, os.path.basename(file_path), metadata)
            
            logger.info(f"Análise de metadados concluída para: {file_path}")
            return metadata
            
        except Exception as e:
            logger.error(f"Erro ao analisar metadados: {str(e)}")
            return {"error": str(e)}
    
    def analyze_image(self, image_path: str) -> Dict[str, Any]:
        """
        Extrai metadados de uma imagem.
        
        Args:
            image_path: Caminho para a imagem
            
        Returns:
            Dicionário com metadados da imagem
        """
        try:
            logger.info(f"Extraindo metadados da imagem: {image_path}")
            
            # Extrair metadados EXIF usando exifread
            with open(image_path, 'rb') as f:
                exif_tags = exifread.process_file(f, details=False)
            
            # Converter tags para dicionário
            exif_data = {}
            for tag, value in exif_tags.items():
                # Ignorar tags de miniatura
                if tag.startswith('Thumbnail'):
                    continue
                
                # Converter valor para string para garantir serialização JSON
                exif_data[tag] = str(value)
            
            # Extrair metadados adicionais usando PIL
            try:
                with Image.open(image_path) as img:
                    # Dimensões da imagem
                    width, height = img.size
                    
                    # Modo de cor
                    mode = img.mode
                    
                    # Formato
                    format_str = img.format
                    
                    # Extrair dados EXIF usando PIL
                    pil_exif = {}
                    if hasattr(img, '_getexif') and img._getexif():
                        exif_info = img._getexif()
                        if exif_info:
                            for tag, value in exif_info.items():
                                tag_name = TAGS.get(tag, tag)
                                # Converter valores não serializáveis para string
                                if isinstance(value, bytes):
                                    value = value.hex()
                                elif isinstance(value, (datetime.datetime, datetime.date, datetime.time)):
                                    value = value.isoformat()
                                
                                pil_exif[tag_name] = value
                    
                    # Extrair dados GPS se disponíveis
                    gps_info = self._extract_gps_info(img)
            except Exception as e:
                logger.warning(f"Erro ao extrair metadados adicionais com PIL: {str(e)}")
                width, height, mode, format_str, pil_exif, gps_info = None, None, None, None, {}, {}
            
            # Compilar resultados
            metadata = {
                'file_type': 'image',
                'dimensions': {
                    'width': width,
                    'height': height
                },
                'color_mode': mode,
                'format': format_str,
                'exif_data': exif_data,
                'pil_exif': pil_exif,
                'gps_info': gps_info,
                'analysis_date': datetime.datetime.now().isoformat()
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Erro ao analisar imagem: {str(e)}")
            return {"error": str(e)}
    
    def analyze_document(self, document_path: str) -> Dict[str, Any]:
        """
        Extrai metadados de um documento Word.
        
        Args:
            document_path: Caminho para o documento
            
        Returns:
            Dicionário com metadados do documento
        """
        try:
            logger.info(f"Extraindo metadados do documento: {document_path}")
            
            # Carregar documento
            doc = Document(document_path)
            
            # Extrair propriedades do documento
            core_properties = doc.core_properties
            
            # Extrair estatísticas do documento
            paragraphs_count = len(doc.paragraphs)
            words_count = sum(len(paragraph.text.split()) for paragraph in doc.paragraphs if paragraph.text)
            
            # Contar tabelas
            tables_count = len(doc.tables)
            
            # Extrair estilos
            styles = [style.name for style in doc.styles]
            
            # Extrair seções
            sections_count = len(doc.sections)
            
            # Compilar resultados
            metadata = {
                'file_type': 'document',
                'core_properties': {
                    'author': core_properties.author,
                    'category': core_properties.category,
                    'comments': core_properties.comments,
                    'content_status': core_properties.content_status,
                    'created': core_properties.created.isoformat() if core_properties.created else None,
                    'identifier': core_properties.identifier,
                    'keywords': core_properties.keywords,
                    'language': core_properties.language,
                    'last_modified_by': core_properties.last_modified_by,
                    'last_printed': core_properties.last_printed.isoformat() if core_properties.last_printed else None,
                    'modified': core_properties.modified.isoformat() if core_properties.modified else None,
                    'revision': core_properties.revision,
                    'subject': core_properties.subject,
                    'title': core_properties.title,
                    'version': core_properties.version
                },
                'statistics': {
                    'paragraphs_count': paragraphs_count,
                    'words_count': words_count,
                    'tables_count': tables_count,
                    'sections_count': sections_count
                },
                'styles': styles,
                'analysis_date': datetime.datetime.now().isoformat()
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Erro ao analisar documento: {str(e)}")
            return {"error": str(e)}
    
    def analyze_spreadsheet(self, spreadsheet_path: str) -> Dict[str, Any]:
        """
        Extrai metadados de uma planilha Excel.
        
        Args:
            spreadsheet_path: Caminho para a planilha
            
        Returns:
            Dicionário com metadados da planilha
        """
        try:
            logger.info(f"Extraindo metadados da planilha: {spreadsheet_path}")
            
            # Verificar se é um arquivo CSV
            if spreadsheet_path.lower().endswith('.csv'):
                return self._analyze_csv(spreadsheet_path)
            
            # Carregar planilha
            wb = load_workbook(spreadsheet_path, read_only=True, data_only=True)
            
            # Extrair propriedades
            properties = wb.properties
            
            # Extrair informações das planilhas
            sheets_info = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                
                # Obter dimensões da planilha
                if sheet.max_row and sheet.max_column:
                    sheet_info = {
                        'name': sheet_name,
                        'rows': sheet.max_row,
                        'columns': sheet.max_column,
                        'cell_count': sheet.max_row * sheet.max_column
                    }
                else:
                    sheet_info = {
                        'name': sheet_name,
                        'rows': 'Desconhecido',
                        'columns': 'Desconhecido',
                        'cell_count': 'Desconhecido'
                    }
                
                sheets_info.append(sheet_info)
            
            # Compilar resultados
            metadata = {
                'file_type': 'spreadsheet',
                'properties': {
                    'creator': properties.creator,
                    'last_modified_by': properties.lastModifiedBy,
                    'created': properties.created.isoformat() if properties.created else None,
                    'modified': properties.modified.isoformat() if properties.modified else None,
                    'title': properties.title,
                    'subject': properties.subject,
                    'description': properties.description,
                    'keywords': properties.keywords,
                    'category': properties.category,
                    'company': properties.company
                },
                'sheets': sheets_info,
                'sheet_count': len(wb.sheetnames),
                'analysis_date': datetime.datetime.now().isoformat()
            }
            
            # Fechar planilha
            wb.close()
            
            return metadata
            
        except Exception as e:
            logger.error(f"Erro ao analisar planilha: {str(e)}")
            return {"error": str(e)}
    
    def analyze_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """
        Extrai metadados de um arquivo PDF.
        
        Args:
            pdf_path: Caminho para o PDF
            
        Returns:
            Dicionário com metadados do PDF
        """
        try:
            logger.info(f"Extraindo metadados do PDF: {pdf_path}")
            
            # Abrir PDF
            with open(pdf_path, 'rb') as f:
                pdf = PdfReader(f)
                
                # Extrair informações do documento
                info = pdf.metadata
                
                # Converter para dicionário
                if info:
                    metadata_dict = {}
                    for key, value in info.items():
                        # Remover o prefixo '/' das chaves
                        clean_key = key[1:] if key.startswith('/') else key
                        
                        # Converter valores para string para garantir serialização JSON
                        if isinstance(value, bytes):
                            value = value.decode('utf-8', errors='replace')
                        
                        metadata_dict[clean_key] = value
                else:
                    metadata_dict = {}
                
                # Extrair estatísticas do documento
                page_count = len(pdf.pages)
                
                # Verificar se o PDF está criptografado
                is_encrypted = pdf.is_encrypted
                
                # Extrair tamanho das páginas
                page_sizes = []
                for i, page in enumerate(pdf.pages):
                    if i < 5:  # Limitar a 5 páginas para performance
                        if '/MediaBox' in page:
                            page_sizes.append(str(page['/MediaBox']))
                
                # Compilar resultados
                metadata = {
                    'file_type': 'pdf',
                    'metadata': metadata_dict,
                    'page_count': page_count,
                    'is_encrypted': is_encrypted,
                    'page_sizes': page_sizes,
                    'analysis_date': datetime.datetime.now().isoformat()
                }
                
                return metadata
            
        except Exception as e:
            logger.error(f"Erro ao analisar PDF: {str(e)}")
            return {"error": str(e)}
    
    def analyze_generic_file(self, file_path: str) -> Dict[str, Any]:
        """
        Extrai metadados básicos de um arquivo genérico.
        
        Args:
            file_path: Caminho para o arquivo
            
        Returns:
            Dicionário com metadados básicos do arquivo
        """
        try:
            logger.info(f"Extraindo metadados básicos do arquivo: {file_path}")
            
            # Obter informações básicas do arquivo
            file_info = self._get_file_info(file_path)
            
            # Tentar determinar o tipo de arquivo pelo conteúdo
            file_type = self._determine_file_type(file_path)
            
            # Compilar resultados
            metadata = {
                'file_type': 'generic',
                'detected_type': file_type,
                'file_info': file_info,
                'analysis_date': datetime.datetime.now().isoformat()
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Erro ao analisar arquivo genérico: {str(e)}")
            return {"error": str(e)}
    
    def analyze_directory(self, directory_path: str) -> Dict[str, Any]:
        """
        Analisa todos os arquivos em um diretório.
        
        Args:
            directory_path: Caminho para o diretório
            
        Returns:
            Dicionário com resumo da análise
        """
        try:
            logger.info(f"Analisando diretório: {directory_path}")
            
            # Verificar se o diretório existe
            if not os.path.exists(directory_path) or not os.path.isdir(directory_path):
                logger.error(f"Diretório não encontrado: {directory_path}")
                return {"error": f"Diretório não encontrado: {directory_path}"}
            
            # Inicializar contadores
            file_count = 0
            analyzed_count = 0
            error_count = 0
            file_types = {}
            
            # Analisar cada arquivo no diretório
            for root, _, files in os.walk(directory_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    file_count += 1
                    
                    try:
                        # Analisar arquivo
                        result = self.analyze_file(file_path)
                        
                        # Verificar se houve erro
                        if 'error' in result:
                            error_count += 1
                        else:
                            analyzed_count += 1
                            
                            # Contar tipos de arquivo
                            file_type = result.get('file_type', 'unknown')
                            if file_type in file_types:
                                file_types[file_type] += 1
                            else:
                                file_types[file_type] = 1
                    
                    except Exception as e:
                        logger.error(f"Erro ao analisar {file_path}: {str(e)}")
                        error_count += 1
            
            # Compilar resultados
            summary = {
                'directory': directory_path,
                'file_count': file_count,
                'analyzed_count': analyzed_count,
                'error_count': error_count,
                'file_types': file_types,
                'analysis_date': datetime.datetime.now().isoformat()
            }
            
            # Salvar resumo
            dir_name = os.path.basename(os.path.normpath(directory_path))
            self._save_results('others', f"directory_{dir_name}", summary)
            
            logger.info(f"Análise de diretório concluída. Analisados {analyzed_count} de {file_count} arquivos.")
            return summary
            
        except Exception as e:
            logger.error(f"Erro ao analisar diretório: {str(e)}")
            return {"error": str(e)}
    
    def export_results_to_csv(self, file_type: str, filename: Optional[str] = None) -> str:
        """
        Exporta os resultados para um arquivo CSV.
        
        Args:
            file_type: Tipo de arquivo ('images', 'documents', 'spreadsheets', 'pdfs', 'others')
            filename: Nome do arquivo (opcional)
            
        Returns:
            Caminho para o arquivo CSV gerado
        """
        try:
            if file_type not in self.results or not self.results[file_type]:
                logger.warning(f"Não há resultados para exportar do tipo: {file_type}")
                return ""
            
            if not filename:
                filename = f"{file_type}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
            
            output_file = os.path.join(self.output_dir, filename)
            
            # Converter resultados para DataFrame
            data = []
            
            for file_name, metadata in self.results[file_type].items():
                # Extrair informações básicas
                row = {
                    'file_name': file_name,
                    'file_type': metadata.get('file_type', ''),
                    'file_size': metadata.get('file_size', ''),
                    'creation_date': metadata.get('creation_date', ''),
                    'modification_date': metadata.get('modification_date', ''),
                    'analysis_date': metadata.get('analysis_date', '')
                }
                
                # Adicionar informações específicas por tipo de arquivo
                if file_type == 'images':
                    if 'dimensions' in metadata:
                        row['width'] = metadata['dimensions'].get('width', '')
                        row['height'] = metadata['dimensions'].get('height', '')
                    
                    if 'gps_info' in metadata and metadata['gps_info']:
                        row['latitude'] = metadata['gps_info'].get('latitude', '')
                        row['longitude'] = metadata['gps_info'].get('longitude', '')
                
                elif file_type == 'documents':
                    if 'core_properties' in metadata:
                        props = metadata['core_properties']
                        row['author'] = props.get('author', '')
                        row['title'] = props.get('title', '')
                        row['created'] = props.get('created', '')
                        row['modified'] = props.get('modified', '')
                    
                    if 'statistics' in metadata:
                        stats = metadata['statistics']
                        row['words_count'] = stats.get('words_count', '')
                        row['paragraphs_count'] = stats.get('paragraphs_count', '')
                
                elif file_type == 'pdfs':
                    if 'metadata' in metadata:
                        pdf_meta = metadata['metadata']
                        row['author'] = pdf_meta.get('Author', '')
                        row['creator'] = pdf_meta.get('Creator', '')
                        row['producer'] = pdf_meta.get('Producer', '')
                        row['title'] = pdf_meta.get('Title', '')
                    
                    row['page_count'] = metadata.get('page_count', '')
                    row['is_encrypted'] = metadata.get('is_encrypted', '')
                
                data.append(row)
            
            # Criar DataFrame
            df = pd.DataFrame(data)
            
            # Salvar como CSV
            df.to_csv(output_file, index=False, encoding='utf-8')
            
            logger.info(f"Resultados exportados para: {output_file}")
            return output_file
            
        except Exception as e:
            logger.error(f"Erro ao exportar resultados para CSV: {str(e)}")
            return ""
    
    def _get_file_info(self, file_path: str) -> Dict[str, Any]:
        """
        Obtém informações básicas de um arquivo.
        
        Args:
            file_path: Caminho para o arquivo
            
        Returns:
            Dicionário com informações básicas do arquivo
        """
        # Obter estatísticas do arquivo
        stat = os.stat(file_path)
        
        # Extrair informações
        file_info = {
            'file_name': os.path.basename(file_path),
            'file_path': file_path,
            'file_size': stat.st_size,
            'file_size_human': self._format_size(stat.st_size),
            'creation_date': datetime.datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modification_date': datetime.datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'access_date': datetime.datetime.fromtimestamp(stat.st_atime).isoformat(),
            'extension': os.path.splitext(file_path)[1].lower()
        }
        
        return file_info
    
    def _determine_file_type(self, file_path: str) -> str:
        """
        Tenta determinar o tipo de arquivo pelo conteúdo.
        
        Args:
            file_path: Caminho para o arquivo
            
        Returns:
            Tipo de arquivo detectado
        """
        # Ler os primeiros bytes do arquivo
        try:
            with open(file_path, 'rb') as f:
                header = f.read(16)
            
            # Verificar assinaturas de arquivo comuns
            if header.startswith(b'\xFF\xD8\xFF'):
                return 'jpeg_image'
            elif header.startswith(b'\x89PNG\r\n\x1A\n'):
                return 'png_image'
            elif header.startswith(b'GIF87a') or header.startswith(b'GIF89a'):
                return 'gif_image'
            elif header.startswith(b'%PDF'):
                return 'pdf'
            elif header.startswith(b'PK\x03\x04'):
                # Pode ser DOCX, XLSX, PPTX, ZIP, etc.
                if file_path.endswith('.docx'):
                    return 'docx_document'
                elif file_path.endswith('.xlsx'):
                    return 'xlsx_spreadsheet'
                elif file_path.endswith('.pptx'):
                    return 'pptx_presentation'
                else:
                    return 'zip_archive'
            elif header.startswith(b'\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1'):
                # Formato de Documento Composto OLE (DOC, XLS, PPT)
                if file_path.endswith('.doc'):
                    return 'doc_document'
                elif file_path.endswith('.xls'):
                    return 'xls_spreadsheet'
                elif file_path.endswith('.ppt'):
                    return 'ppt_presentation'
                else:
                    return 'ole_compound_document'
            else:
                # Verificar se é um arquivo de texto
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        f.read(1024)
                    return 'text'
                except UnicodeDecodeError:
                    return 'binary'
        
        except Exception as e:
            logger.warning(f"Erro ao determinar tipo de arquivo: {str(e)}")
            return 'unknown'
    
    def _extract_gps_info(self, img: Image.Image) -> Dict[str, Any]:
        """
        Extrai informações GPS de uma imagem.
        
        Args:
            img: Objeto PIL Image
            
        Returns:
            Dicionário com informações GPS
        """
        gps_info = {}
        
        try:
            if not hasattr(img, '_getexif') or not img._getexif():
                return gps_info
            
            exif_info = img._getexif()
            if not exif_info:
                return gps_info
            
            # Procurar pela tag GPS
            gps_tag = None
            for tag, value in TAGS.items():
                if value == 'GPSInfo':
                    gps_tag = tag
                    break
            
            if gps_tag and gps_tag in exif_info:
                gps_data = exif_info[gps_tag]
                
                # Extrair informações GPS
                for tag, value in gps_data.items():
                    tag_name = GPSTAGS.get(tag, tag)
                    gps_info[tag_name] = value
                
                # Calcular coordenadas em formato decimal
                if 'GPSLatitude' in gps_info and 'GPSLatitudeRef' in gps_info:
                    lat = self._convert_to_degrees(gps_info['GPSLatitude'])
                    if gps_info['GPSLatitudeRef'] == 'S':
                        lat = -lat
                    gps_info['latitude'] = lat
                
                if 'GPSLongitude' in gps_info and 'GPSLongitudeRef' in gps_info:
                    lon = self._convert_to_degrees(gps_info['GPSLongitude'])
                    if gps_info['GPSLongitudeRef'] == 'W':
                        lon = -lon
                    gps_info['longitude'] = lon
                
                # Adicionar link para Google Maps se tiver coordenadas
                if 'latitude' in gps_info and 'longitude' in gps_info:
                    gps_info['google_maps_url'] = f"https://maps.google.com/maps?q={gps_info['latitude']},{gps_info['longitude']}"
        
        except Exception as e:
            logger.warning(f"Erro ao extrair informações GPS: {str(e)}")
        
        return gps_info
    
    def _convert_to_degrees(self, value: tuple) -> float:
        """
        Converte coordenadas GPS de graus, minutos, segundos para formato decimal.
        
        Args:
            value: Tupla com valores de graus, minutos, segundos
            
        Returns:
            Coordenada em formato decimal
        """
        d, m, s = value
        
        # Converter para valores numéricos
        if isinstance(d, tuple):
            d = d[0] / d[1]
        
        if isinstance(m, tuple):
            m = m[0] / m[1]
        
        if isinstance(s, tuple):
            s = s[0] / s[1]
        
        return d + (m / 60.0) + (s / 3600.0)
    
    def _analyze_csv(self, csv_path: str) -> Dict[str, Any]:
        """
        Analisa um arquivo CSV.
        
        Args:
            csv_path: Caminho para o arquivo CSV
            
        Returns:
            Dicionário com metadados do CSV
        """
        try:
            # Tentar diferentes delimitadores
            for delimiter in [',', ';', '\t', '|']:
                try:
                    df = pd.read_csv(csv_path, delimiter=delimiter, nrows=5)
                    if len(df.columns) > 1:
                        break
                except:
                    continue
            
            # Obter estatísticas
            row_count = sum(1 for _ in open(csv_path, 'r', encoding='utf-8', errors='replace'))
            column_count = len(df.columns)
            
            # Compilar resultados
            metadata = {
                'file_type': 'csv',
                'delimiter': delimiter,
                'row_count': row_count,
                'column_count': column_count,
                'columns': df.columns.tolist(),
                'sample_data': df.head(5).to_dict(orient='records'),
                'analysis_date': datetime.datetime.now().isoformat()
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Erro ao analisar CSV: {str(e)}")
            return {
                'file_type': 'csv',
                'error': str(e),
                'analysis_date': datetime.datetime.now().isoformat()
            }
    
    def _format_size(self, size_bytes: int) -> str:
        """
        Formata o tamanho do arquivo para formato legível.
        
        Args:
            size_bytes: Tamanho em bytes
            
        Returns:
            Tamanho formatado
        """
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.2f} {unit}"
            size_bytes /= 1024.0
        
        return f"{size_bytes:.2f} PB"
    
    def _save_results(self, file_type: str, identifier: str, data: Dict[str, Any]):
        """
        Salva os resultados em um arquivo JSON.
        
        Args:
            file_type: Tipo de arquivo ('images', 'documents', 'spreadsheets', 'pdfs', 'others')
            identifier: Identificador único para os resultados
            data: Dados a serem salvos
        """
        output_file = os.path.join(
            self.output_dir, 
            file_type, 
            f"{identifier}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        )
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
        
        logger.info(f"Resultados salvos em: {output_file}")
        
        # Armazenar resultados na memória
        if file_type not in self.results:
            self.results[file_type] = {}
        
        self.results[file_type][identifier] = data


# Função para demonstração do módulo
def demo():
    """Função de demonstração do módulo de análise de metadados."""
    osint = MetadataAnalysisOSINT(output_dir="resultados_osint")
    
    # Criar um arquivo de texto para teste
    test_file_path = os.path.join(osint.output_dir, "test_document.txt")
    with open(test_file_path, 'w', encoding='utf-8') as f:
        f.write("Este é um arquivo de teste para demonstração do módulo de análise de metadados.\n")
        f.write("Ele contém informações básicas para extração de metadados.\n")
        f.write("OSINT - Open Source Intelligence\n")
    
    # Demonstração de análise de arquivo
    print("Demonstração de análise de arquivo:")
    file_results = osint.analyze_file(test_file_path)
    print(f"Arquivo: {file_results['file_name']}")
    print(f"Tamanho: {file_results['file_size_human']}")
    print(f"Data de criação: {file_results['creation_date']}")
    
    # Demonstração de análise de diretório
    print("\nDemonstração de análise de diretório:")
    dir_results = osint.analyze_directory(osint.output_dir)
    print(f"Diretório: {dir_results['directory']}")
    print(f"Arquivos analisados: {dir_results['analyzed_count']} de {dir_results['file_count']}")
    
    # Exportar resultados
    csv_file = osint.export_results_to_csv('others')
    if csv_file:
        print(f"\nResultados exportados para: {csv_file}")
    
    print("\nDemonstração concluída!")


if __name__ == "__main__":
    demo()
